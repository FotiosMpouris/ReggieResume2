import re
from docx import Document
from io import BytesIO
from fpdf import FPDF
import logging

class ReggieResume:
    def __init__(self):
        self.personal_info = {}
        self.education = []
        self.experience = []
        self.skills = {}

    # Function to directly input entire resume as plain text
    def input_resume(self, resume_text):
        logging.info("Storing the provided resume.")
        self.resume_text = resume_text

    # Function to extract skills from the resume based on predefined patterns
    def extract_skills(self):
        logging.info("Extracting skills from the resume.")
        # Example: you can define skill patterns to search within the resume text
        skills_pattern = re.compile(r"\b(skill[s]?|proficienc[yi]?|technolog[yi]?|languages?)\b", re.I)
        matched_skills = re.findall(skills_pattern, self.resume_text)
        return list(set(matched_skills))  # Removing duplicates

    # Function to compare job description with resume skills
    def compare_skills(self, job_description):
        logging.info("Comparing skills from job description with resume.")
        job_skills = re.findall(r"\b(?:skill[s]?|expertise|knowledge)\b.*?:?\s*(.*?)\s*[.,;]", job_description, re.I)
        resume_skills = self.extract_skills()

        # Compare and create a dictionary of matched vs unmatched skills
        matched_skills = [skill for skill in job_skills if skill.lower() in map(str.lower, resume_skills)]
        unmatched_skills = [skill for skill in job_skills if skill.lower() not in map(str.lower, resume_skills)]

        return matched_skills, unmatched_skills

    # Function to generate a tailored resume based on matched skills
    def generate_resume(self, job_description):
        logging.info("Generating tailored resume based on job description.")
        matched_skills, unmatched_skills = self.compare_skills(job_description)

        tailored_resume = f"--- Tailored Resume ---\n\n{self.resume_text}\n\n"
        tailored_resume += "--- Skills Comparison ---\n"
        tailored_resume += f"\nMatched Skills from Job Description: {', '.join(matched_skills)}\n"
        tailored_resume += f"Unmatched Skills: {', '.join(unmatched_skills)}\n"

        # Generate a downloadable DOCX
        docx_buffer = BytesIO()
        doc = Document()
        doc.add_heading('Tailored Resume', 0)
        doc.add_paragraph(self.resume_text)
        doc.add_heading('Skills Comparison', level=1)
        doc.add_paragraph(f"Matched Skills: {', '.join(matched_skills)}")
        doc.add_paragraph(f"Unmatched Skills: {', '.join(unmatched_skills)}")
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return tailored_resume, docx_buffer

    # Function to generate PDF from text
    def generate_pdf(self, text):
        logging.info("Generating PDF version of the resume.")
        pdf_buffer = BytesIO()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in text.split("\n"):
            pdf.cell(200, 10, txt=line, ln=True)

        pdf.output(pdf_buffer)
        pdf_buffer.seek(0)
        return pdf_buffer

# The class can now handle full resume text input and job comparison directly.
